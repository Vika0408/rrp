import spire.doc
from flask import Flask, render_template, request, send_file, flash, url_for, after_this_request, redirect
import os
import uuid
import docx
from datetime import datetime
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import pandas as pd
from pathlib2 import Path
from spire.doc import *
from spire.doc.common import *
import re

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/process', methods=['POST'])
def process():
    if 'folder' not in request.files or not request.files.getlist('folder'):
        return redirect(url_for('index'))

    folders = request.files.getlist('folder')

    unique_folder_name = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(uuid.uuid4())
    upload_folder_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_folder_name)

    Path(upload_folder_path).mkdir(parents=True, exist_ok=True)

    data = []

    for file in folders:
        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(upload_folder_path, filename)
            file.save(file_path)

            text = ""
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx') or filename.endswith('.doc'):
                if filename.endswith('.doc'):
                    convert_doc_to_docx(file_path)
                    filename = f"{os.path.splitext(filename)[0]}.docx"  # Update filename after conversion
                text = extract_text_from_docx(os.path.join(upload_folder_path, filename))  # Use updated filename

            emails, phone_numbers, cleaned_text = extract_information_from_text(text)
            data.append({'Filename': filename, 'Text': cleaned_text, 'Emails': emails, 'Phone Numbers': phone_numbers})

    df = pd.DataFrame(data)

    excel_filename = os.path.join(upload_folder_path, 'parsed_data.xlsx')
    df.to_excel(excel_filename, index=False)

    return send_file(excel_filename, as_attachment=True)


def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
    return text


def extract_text_from_docx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def convert_doc_to_docx(file_path):
    document = spire.doc.Document()
    document.LoadFromFile(file_path)
    document.SaveToFile(f"{os.path.splitext(file_path)[0]}.docx", FileFormat.Docx2016)


def extract_information_from_text(text):
    # Regular expressions to match email addresses and phone numbers
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(?:(?:\+?(\d{1,3}))?[-. ]?)?(\d{3})[-. ]?(\d{3})[-. ]?(\d{4})'

    # Find all email addresses in the text
    emails = re.findall(email_pattern, text)

    # Find all phone numbers in the text
    phone_numbers = re.findall(phone_pattern, text)
    formatted_phone_numbers = [''.join(phone) for phone in phone_numbers]

    # Remove duplicates and convert to set for unique values
    unique_emails = list(set(emails))
    unique_phone_numbers = list(set(formatted_phone_numbers))

    # Remove empty strings from the list of phone numbers
    unique_phone_numbers = [phone for phone in unique_phone_numbers if phone]

    # Remove email addresses and phone numbers from the text
    cleaned_text = re.sub(email_pattern, '', text)
    cleaned_text = re.sub(phone_pattern, '', cleaned_text)

    return unique_emails, unique_phone_numbers, cleaned_text


if __name__ == '__main__':
    app.run(debug=True)
